#!/usr/bin/env python
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import logging
import os
import sys
import traceback
from datetime import datetime
from pathlib import Path
from typing import Iterator, List, Dict, Optional

def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Ingesta de Documentos — TFM RAG")
    p.add_argument("--input-dir", required=True, help="Carpeta base a ingerir")
    p.add_argument("--pattern", action="append", default=None,
                   help="Patrones glob separados por coma (p.ej. *.pdf,*.docx). Repetible.")
    p.add_argument("--recursive", dest="recursive", action="store_true", default=True)
    p.add_argument("--no-recursive", dest="recursive", action="store_false")
    p.add_argument("--only-new", action="store_true", default=False, help="Procesar solo nuevos/modificados")

    # Legado / compat
    p.add_argument("--include-ext", nargs="+", default=None)
    p.add_argument("--policy", choices=["hash", "mtime"], default="hash")

    # CSV / texto
    p.add_argument("--encoding-default", default="utf-8")
    p.add_argument("--csv-delimiter", default=",")
    p.add_argument("--csv-quotechar", default='"')
    p.add_argument("--csv-header", dest="csv_header", action="store_true", default=True)
    p.add_argument("--no-csv-header", dest="csv_header", action="store_false")
    p.add_argument("--csv-columns", nargs="*", default=None)

    # Chunking
    p.add_argument("--chunk-size", type=int, default=512)
    p.add_argument("--chunk-overlap", type=int, default=64)

    # Salidas
    p.add_argument("--verbose-json", action="store_true")
    p.add_argument("--run-dir", default=None, help="Directorio donde guardar artefactos de este run")

    # Proyecto
    p.add_argument("--project-root", default=".", help="Ruta del proyecto (donde está app/)")
    return p

def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()

def _fingerprint(path: Path, policy: str) -> str:
    if policy == "hash":
        return _sha256_file(path)
    st = path.stat()
    return f"{st.st_size}:{int(st.st_mtime_ns)}"

def _split_patterns(patterns_arg: Optional[List[str]], include_ext: Optional[List[str]]) -> List[str]:
    if patterns_arg:
        pats: List[str] = []
        for entry in patterns_arg:
            pats += [p.strip() for p in entry.split(",") if p.strip()]
        return pats
    if include_ext:
        return [f"*.{e.lstrip('.')}" for e in include_ext]
    return ["*.pdf", "*.docx", "*.txt", "*.md", "*.csv"]

def _iter_files(base: Path, patterns: List[str], recursive: bool) -> Iterator[Path]:
    if recursive:
        for pat in patterns:
            yield from base.rglob(pat)
    else:
        for pat in patterns:
            yield from base.glob(pat)

def _read_text_file(path: Path, enc: str) -> str:
    return path.read_text(encoding=enc, errors="ignore")

def _read_csv(path: Path, delimiter: str, quotechar: str, header: bool, columns: Optional[List[str]], enc: str) -> str:
    out_rows: List[str] = []
    with path.open("r", encoding=enc, errors="ignore", newline="") as f:
        reader = csv.reader(f, delimiter=delimiter, quotechar=quotechar)
        cols_idx: Optional[List[int]] = None
        for i, row in enumerate(reader):
            if i == 0 and header:
                headers = row
                if columns:
                    cols_idx = [headers.index(c) for c in columns if c in headers]
                continue
            if columns and cols_idx is not None:
                row = [row[j] for j in cols_idx if j < len(row)]
            out_rows.append(" ".join(row))
    return "\n".join(out_rows)

def _read_pdf(path: Path) -> Optional[str]:
    try:
        import PyPDF2  # type: ignore
    except Exception:
        return None
    try:
        txt_parts: List[str] = []
        with path.open("rb") as f:
            pdf = PyPDF2.PdfReader(f)
            for page in pdf.pages:
                txt_parts.append(page.extract_text() or "")
        return "\n".join(txt_parts).strip() or None
    except Exception:
        return None

def _read_docx(path: Path) -> Optional[str]:
    try:
        import docx  # type: ignore
    except Exception:
        return None
    try:
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs).strip() or None
    except Exception:
        return None

def _chunk_text(text: str, size: int, overlap: int) -> List[str]:
    if size <= 0:
        return [text]
    chunks: List[str] = []
    i = 0
    n = len(text)
    step = max(1, size - max(0, overlap))
    while i < n:
        chunks.append(text[i:i + size])
        i += step
    return chunks

def _setup_logging() -> None:
    logs_dir = Path("data/logs"); logs_dir.mkdir(parents=True, exist_ok=True)
    log_path = logs_dir / "ingestion.log"
    fmt = logging.Formatter("[%(asctime)s] %(levelname)s %(message)s")
    root = logging.getLogger()
    root.handlers.clear()
    root.setLevel(logging.INFO)
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setFormatter(fmt)
    sh = logging.StreamHandler(sys.stdout)
    sh.setFormatter(fmt)
    root.addHandler(fh)
    root.addHandler(sh)

def main() -> int:
    args = build_parser().parse_args()

    project_root = Path(args.project_root).resolve()
    if not (project_root / "app").exists():
        sys.stderr.write(f"No encuentro la carpeta 'app' bajo {project_root}.\n")
        return 2

    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))
    os.chdir(project_root)

    _setup_logging()
    logging.info("== Ingesta DOCS iniciada ==")

    try:
        from app import create_app  # type: ignore
        from app.extensions.db import get_session  # type: ignore
        from app.models import Source, Document, Chunk  # type: ignore
    except Exception as e:
        logging.exception("ImportError")
        sys.stderr.write(f"ImportError: {e}\n")
        return 2

    app = create_app()
    logging.info("App creada: Prototipo_chatbot")

    base = Path(args.input_dir).resolve()
    base.mkdir(parents=True, exist_ok=True)

    patterns = _split_patterns(args.pattern, args.include_ext)
    manifest_path = Path("data/processed/documents/manifest.json")
    if manifest_path.exists():
        try:
            manifest: Dict[str, Dict[str, str]] = json.loads(manifest_path.read_text(encoding="utf-8"))
        except Exception:
            manifest = {}
    else:
        manifest = {}

    # run_dir
    if args.run_dir:
        run_dir = Path(args.run_dir)
    else:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        run_dir = Path("data/processed/runs/docs") / f"run_{ts}"
    run_dir.mkdir(parents=True, exist_ok=True)
    logging.info("run_dir: %s", run_dir)

    stats = {
        "scanned": 0,
        "new_docs": 0,
        "updated_docs": 0,
        "skipped_unchanged": 0,
        "failed": 0,
        "total_chunks": 0,
    }

    from time import perf_counter
    t0 = perf_counter()

    # Asegurar Source (por URL base)
    with get_session() as s:
        src = s.query(Source).filter(Source.type == "docs", Source.url == str(base)).first()
        if not src:
            src = Source(type="docs", url=str(base), name=base.name, config={
                "input_dir": str(base),
                "patterns": patterns,
                "recursive": bool(args.recursive),
                "policy": args.policy,
            })
            s.add(src)
            s.flush()
        source_id = src.id

    files = list(_iter_files(base, patterns, args.recursive))
    files = sorted({f.resolve() for f in files if f.is_file()})
    logging.info("Enumerados %d ficheros", len(files))

    for path in files:
        stats["scanned"] += 1
        try:
            fp = _fingerprint(path, args.policy)
            key = str(path)
            prev = manifest.get(key)
            unchanged = bool(prev and prev.get("fp") == fp)

            if args.only_new and unchanged:
                stats["skipped_unchanged"] += 1
                logging.info("SKIP (unchanged, only_new): %s", path)
                continue

            # leer contenido (si podemos)
            ext = path.suffix.lower().lstrip(".")
            content: Optional[str] = None
            if ext in ("txt", "md"):
                content = _read_text_file(path, args.encoding_default)
            elif ext == "csv":
                content = _read_csv(path, args.csv_delimiter, args.csv_quotechar, args.csv_header, args.csv_columns, args.encoding_default)
            elif ext == "pdf":
                content = _read_pdf(path)
            elif ext == "docx":
                content = _read_docx(path)

            rechunked = False

            with get_session() as s:
                doc = s.query(Document).filter(
                    Document.source_id == source_id, Document.path == str(path)
                ).first()
                if doc is None:
                    st = path.stat()
                    doc = Document(
                        source_id=source_id,
                        path=str(path),
                        title=path.name,
                        ext=ext,
                        size=st.st_size,
                        mtime_ns=int(st.st_mtime_ns),
                        hash=fp if args.policy == "hash" else None,
                        meta={"policy": args.policy},
                    )
                    s.add(doc)
                    s.flush()
                    created = True
                else:
                    st = path.stat()
                    doc.size = st.st_size
                    doc.mtime_ns = int(st.st_mtime_ns)
                    if args.policy == "hash":
                        doc.hash = fp
                    created = False

                if content is not None:
                    s.query(Chunk).filter(Chunk.document_id == doc.id).delete()
                    pieces = _chunk_text(content, args.chunk_size, args.chunk_overlap)
                    for i, piece in enumerate(pieces, start=1):
                        s.add(Chunk(
                            source_id=source_id,
                            document_id=doc.id,
                            ordinal=i,
                            text=piece,
                            content=piece,
                            meta={"path": str(path)},
                        ))
                    stats["total_chunks"] += len(pieces)
                    rechunked = True

                s.commit()

            manifest[key] = {"fp": fp, "ts": datetime.now().isoformat()}

            if created:
                stats["new_docs"] += 1
                logging.info("NEW: %s", path)
            else:
                if not unchanged:
                    stats["updated_docs"] += 1
                    logging.info("UPDATED (changed fp): %s", path)
                else:
                    if rechunked:
                        stats["updated_docs"] += 1
                        logging.info("UPDATED (rechunk unchanged): %s", path)
                    else:
                        stats["skipped_unchanged"] += 1
                        logging.info("SKIP (unchanged, no-op): %s", path)

        except Exception:
            stats["failed"] += 1
            logging.exception("Fallo procesando %s", path)
            traceback.print_exc(file=sys.stderr)

    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    elapsed = perf_counter() - t0
    payload = {
        "status": "done",
        "stats": stats,
        "source_dir": str(base),
        "patterns": patterns,
        "policy": args.policy,
        "recursive": bool(args.recursive),
        "only_new": bool(args.only_new),
        "elapsed_sec": round(elapsed, 3),
        "run_dir": str(run_dir.resolve()),
    }

    if args.verbose_json:
        print(json.dumps(payload, ensure_ascii=False))
    else:
        st = {**{k: 0 for k in stats}, **stats}
        print("Status: done")
        print(
            "Stats : scanned={scanned} new={new_docs} updated={updated_docs} "
            "skipped={skipped_unchanged} failed={failed} chunks={total_chunks}".format(**st)
        )
        print(f"Elapsed: {round(elapsed,2)}s")
        print(f"run_dir: {run_dir.resolve()}")

    run_dir.mkdir(parents=True, exist_ok=True)
    (run_dir / "summary.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    out_dir = Path("data/processed/documents/runs"); out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    (out_dir / f"run_docs_{ts}.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    logging.info("== Ingesta DOCS finalizada: %s ==", stats)
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
